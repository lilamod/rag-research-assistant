"""
Extracts raw text from uploaded documents.
Supports: .pdf, .docx, .txt, .md
"""
from pathlib import Path
from pypdf import PdfReader
import docx


class UnsupportedFileTypeError(Exception):
    pass


def load_text(file_path: Path) -> str:
    """Return the full extracted text of a document, given its path."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path)
    elif suffix == ".docx":
        return _load_docx(file_path)
    elif suffix in (".txt", ".md"):
        return _load_plain_text(file_path)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix}'. Supported: .pdf, .docx, .txt, .md"
        )


def _load_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(pages)


def _load_docx(file_path: Path) -> str:
    document = docx.Document(str(file_path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Also pull text out of any tables
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _load_plain_text(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")
